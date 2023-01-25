import docx
import os

path = os.getcwd()
resource_path = os.path.join(path, 'automate_online-materials')

doc = docx.Document(os.path.join(path,'invitation.docx'))

name_list = open(os.path.join(resource_path,'guests.txt'),'r')
names = name_list.readlines()
print(names)

for name in names:
    intro = doc.add_paragraph()
    Intro = intro.add_run('It would be a pleasure to have the company of')

    intro.alignment = 1
    Intro.font.name = 'MathJax_Typewriter'
    doc.add_paragraph('\n')
    Intro.font.size = docx.shared.Pt(20)
    

    Name = intro.add_run('\n\n' + name)
    Name.bold = True
    Name.alignment = 1
    Name.font.name = 'Z003'
    Name.font.size = docx.shared.Pt(40)

    address = doc.add_paragraph()
    Address = address.add_run('at 11010 Memory Lane on the Evening of')
    address.style = 'Normal'
    address.alignment = 1
    Address.font.name = 'MathJax_Typewriter'
    Address.font.size = docx.shared.Pt(20)

    Date = doc.add_paragraph()
    date = Date.add_run('April 1st')
    Date.style = 'Normal'
    Date.alignment = 1
    date.font.name = 'MathJax_Typewriter'
    date.font.size = docx.shared.Pt(20)


    time = doc.add_paragraph()
    Time = time.add_run('at 7 o\'clock')
    time.style = 'Normal'
    time.alignment = 1
    Time.font.name = 'MathJax_Typewriter'
    Time.font.size = docx.shared.Pt(20)


    doc.add_page_break()
    
    

doc.save('invitations.docx')